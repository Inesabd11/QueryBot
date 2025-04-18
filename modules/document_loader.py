import os
import sys
from typing import List
from abc import ABC, abstractmethod
import logging
# files format imports
import PyPDF2 
from docx import Document as DocxDocument
import pandas as pd 
#langchain
from langchain.schema import Document

# Ajouter le chemin racine pour accéder aux modules config
sys.path.append(os.path.abspath(os.path.join(os.path.dirname(__file__), '..')))

from config.paths import DATA_DIR

#logging
logging.basicConfig(level=logging.INFO)
logger = logging.getLogger(__name__)

# --------------------------------------------------------------------
# STEP 1: Abstract Base Class
# --------------------------------------------------------------------

class DocumentLoader(ABC): 
    @abstractmethod
    def load(self, file_path: str) -> List[Document]:
        pass
# --------------------------------------------------------------------
# STEP 2: Different loaders format
# --------------------------------------------------------------------

#PDF 
class PDFLoader(DocumentLoader):
    def load (self, file_path: str) -> List[Document]:
        if not os.path.exists(file_path):
            raise FileNotFoundError(f" Le fichier {file_path} n'existe pas.")
        
        text = ""
        with open (file_path, 'rb') as file: 
            reader = PyPDF2.PdfReader(file)
            for page in reader.pages: 
                page_text = page.extract_text()
                if page_text:
                    text += page_text + "\n"
        logger.info(f"✅ PDF chargé: {file_path}")
        
        return [Document(page_content=text, metadata={"source": file_path})]

#DOCX
class DOCXLoader(DocumentLoader):
    def load(self, file_path: str) -> List[Document]:
        if not os.path.exists(file_path):
            raise FileNotFoundError(f" Le fichier {file_path} n'existe pas.")
        doc = DocxDocument(file_path)
        text = "\n".join([p.text for p in doc.paragraphs if p.text.strip()])
        
        logger.info(f"✅ DOCX chargé: {file_path}")
        return [Document(page_content=text, metadata={"source": file_path})]
    
#CSV
class CSVLoader(DocumentLoader):
    def load(self, file_path: str) -> List[Document]:
        if not os.path.exists(file_path):
            raise FileNotFoundError(f" Le fichier {file_path} n'existe pas.")

        df = pd.read_csv(file_path)
        rows = df.astype(str).values.tolist()
        documents = []
        
        for row in rows:
            text = " | ".join(row)
            documents.append(Document(page_content=text, metadata={"source": file_path}))
        #row_texts = [" | ".join(row) for row in rows]

        logger.info(f"✅ CSV chargé: {file_path}")
        return documents

#TXT
class TXTLoader(DocumentLoader):
    def load(self, file_path: str) -> List[Document]:
        if not os.path.exists(file_path):
            raise FileNotFoundError(f" Le fichier {file_path} n'existe pas.")

        with open(file_path, 'r', encoding='utf-8') as file:
            text = file.read()

        logger.info(f"✅ TXT chargé: {file_path}")
        return [Document(page_content=text, metadata={"source": file_path})]
    
# --------------------------------------------------------------------
# STEP 3: Factory/File Extension
# --------------------------------------------------------------------

class DocumentLoaderFactory:

    @staticmethod
    def get_loader(file_path: str) -> DocumentLoader:
        if not os.path.exists(file_path):
            raise FileNotFoundError(f"Le fichier {file_path} n'existe pas.")

        extension = os.path.splitext(file_path)[1].lower()

        if extension == '.pdf':
            return PDFLoader()
        elif extension == '.docx':
            return DOCXLoader()
        elif extension == '.csv':
            return CSVLoader()
        elif extension == '.txt':
            return TXTLoader()
        else:
            raise ValueError(f"❌ Le format {extension} n'est pas supporté.")

# --------------------------------------------------------------------
# STEP 4: Helper function
# --------------------------------------------------------------------

def load_document(file_path: str) -> List[Document]:
    
    """
    Loads a document from a given file path using the appropriate loader.

    Args:
        file_path (str): The path to the file to be loaded.

    Returns:
        Union[str, List[str]]: The extracted content from the file.
    """
    
    loader = DocumentLoaderFactory.get_loader(file_path)
    return loader.load(file_path)

# --------------------------------------------------------------------
# STEP 5: Load All Files from a Directory
# --------------------------------------------------------------------
def load_all_documents_from_directory(directory: str) -> List[Document]:
    if not os.path.exists(directory):
        raise FileNotFoundError(f"Le dossier {directory} n'existe pas.")

    all_documents = []
    for filename in os.listdir(directory):
        file_path = os.path.join(directory, filename)

        if os.path.isfile(file_path):
            try:
                documents = load_document(file_path)
                all_documents.extend(documents)
            except Exception as e:
                logger.warning(f"⚠️ Erreur lors du chargement de {file_path}: {e}")

    logger.info(f"Total documents chargés: {len(all_documents)}")
    return all_documents

# --------------------------------------------------------------------
# STEP 6: Exemple d’utilisation directe (facultatif)
# --------------------------------------------------------------------

if __name__ == "__main__":
    try:
        content = load_all_documents_from_directory(os.path.join(DATA_DIR, "raw"))
        for doc in content:
            print(doc.page_content[:300])
            print("-" * 50)
    except Exception as e:
        logger.error(f"⚠️ Erreur: {e}")