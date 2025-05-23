import os
import sys
from typing import List
from abc import ABC, abstractmethod
import logging
from pathlib import Path

# Add path to root directory for imports
sys.path.append(os.path.abspath(os.path.join(os.path.dirname(__file__), '..')))

# Import config first
from config.constants import SUPPORTED_FILE_TYPES, validate_file, ALL_SUPPORTED_EXTENSIONS
from config.paths import DATA_DIR, TESSERACT_PATH

# Document processing imports
from PIL import Image
import pytesseract
import PyPDF2
from docx import Document as DocxDocument
import openpyxl
import textract
import pandas as pd
import json
# Technical file formats
import ezdxf  # For CAD files
from stl import mesh as numpy_stl  # Rename to avoid confusion
import xml.etree.ElementTree as ET  # For XML
import yaml  # For YAML/YML
import svglib.svglib  # For SVG
from striprtf.striprtf import rtf_to_text  # For RTF

# Langchain imports
from langchain.schema import Document

# Configure logging
logging.basicConfig(level=logging.INFO)
logger = logging.getLogger(__name__)

# Configure Tesseract path
pytesseract.pytesseract.tesseract_cmd = str(TESSERACT_PATH)

# Verify Tesseract configuration
def verify_tesseract():
    try:
        version = pytesseract.get_tesseract_version()
        logger.info(f"✅ Tesseract configured successfully (version {version})")
        return True
    except Exception as e:
        logger.error(f"❌ Tesseract configuration failed: {e}")
        return False

# Add this check after imports
if not verify_tesseract():
    raise RuntimeError(
        "Tesseract OCR is not properly configured. "
        f"Expected at: {TESSERACT_PATH}"
    )

# --------------------------------------------------------------------
# STEP 1: Abstract Base Class
# --------------------------------------------------------------------

class DocumentLoader(ABC): 
    """Abstract base class for document loaders"""
    @abstractmethod
    def load(self, file_path: str) -> List[Document]:
        pass
# --------------------------------------------------------------------
# STEP 2: Different loaders format
# --------------------------------------------------------------------

#PDF 
class PDFLoader(DocumentLoader):
    def load (self, file_path: str) -> List[Document]:
        if not os.path.exists(file_path):
            raise FileNotFoundError(f" Le fichier {file_path} n'existe pas.")
        
        text = ""
        with open (file_path, 'rb') as file: 
            reader = PyPDF2.PdfReader(file)
            for page in reader.pages: 
                page_text = page.extract_text()
                if page_text:
                    text += page_text + "\n"
        logger.info(f"✅ PDF chargé: {file_path}")
        
        return [Document(page_content=text, metadata={"source": file_path})]
    
# PNG, JPG, JPEG    
class ImageLoader(DocumentLoader):
    """Unified image loader for image formats"""
    def load(self, file_path: str) -> List[Document]:
        file_path = Path(file_path)
        
        try:
            # Validate file exists
            if not file_path.exists():
                raise FileNotFoundError(f"Le fichier {file_path} n'existe pas.")

            # Validate file extension
            if file_path.suffix.lower() not in SUPPORTED_FILE_TYPES['images']:
                raise ValueError(f"Format non supporté: {file_path.suffix}")

            # Load and process image
            with Image.open(file_path) as image:
                # Convert image to RGB if necessary
                if image.mode != 'RGB':
                    image = image.convert('RGB')
                
                # Extract text using OCR
                text = pytesseract.image_to_string(image, lang='fra+eng')
                
                if not text.strip():
                    logger.warning(f"⚠️ Aucun texte extrait de l'image: {file_path}")
                    text = "No text could be extracted from this image."
                
                logger.info(f"✅ Image traitée: {file_path}")
                return [Document(
                    page_content=text,
                    metadata={
                        "source": str(file_path),
                        "format": file_path.suffix.upper().replace('.', ''),
                        "ocr_processed": True,
                        "image_size": f"{image.size[0]}x{image.size[1]}"
                    }
                )]

        except Exception as e:
            logger.error(f"❌ Erreur lors du traitement de l'image {file_path}: {str(e)}")
            raise
    
#xls & xlsx
class ExcelLoader(DocumentLoader):
    def load(self, file_path: str) -> List[Document]:
        if not os.path.exists(file_path):
            raise FileNotFoundError(f"Le fichier {file_path} n'existe pas.")

        # Get file extension for logging
        _, extension = os.path.splitext(file_path)
        extension = extension.lower()
        file_type = extension[1:].upper()  # Remove the dot and convert to uppercase
        
        # Load the workbook - openpyxl works for both .xlsx and newer .xls files
        try:
            wb = openpyxl.load_workbook(file_path)
            sheet = wb.active
            text = ""

            for row in sheet.iter_rows(values_only=True):
                text += " | ".join([str(cell) for cell in row if cell is not None]) + "\n"
                
        except Exception as e:
            # If openpyxl fails for older .xls files, you might need to use xlrd
            if extension == '.xls':
                import xlrd
                wb = xlrd.open_workbook(file_path)
                sheet = wb.sheet_by_index(0)
                text = ""
                
                for row_idx in range(sheet.nrows):
                    row_values = [str(sheet.cell_value(row_idx, col_idx)) for col_idx in range(sheet.ncols) 
                                 if sheet.cell_value(row_idx, col_idx) is not None]
                    text += " | ".join(row_values) + "\n"
            else:
                # If it's not an .xls file, re-raise the exception
                raise e

        logger.info(f"✅ {file_type} chargé: {file_path}")
        return [Document(page_content=text, metadata={"source": file_path})]

#JSON
class JSONLoader(DocumentLoader):
    def load(self, file_path: str) -> List[Document]:
        if not os.path.exists(file_path):
            raise FileNotFoundError(f" Le fichier {file_path} n'existe pas.")

        with open(file_path, 'r', encoding='utf-8') as file:
            data = json.load(file)

        text = json.dumps(data, ensure_ascii=False)

        logger.info(f"✅ JSON chargé: {file_path}")
        return [Document(page_content=text, metadata={"source": file_path})]


#DOCX & DOC
class WordDocumentLoader(DocumentLoader):
    def load(self, file_path: str) -> List[Document]:
        if not os.path.exists(file_path):
            raise FileNotFoundError(f"Le fichier {file_path} n'existe pas.")
        
        # Get file extension
        _, extension = os.path.splitext(file_path)
        extension = extension.lower()
        
        # Process based on file extension
        if extension == '.docx':
            # DOCX handling with python-docx
            doc = DocxDocument(file_path)
            text = "\n".join([p.text for p in doc.paragraphs if p.text.strip()])
            file_type = "DOCX"
        elif extension == '.doc':
            # DOC handling with textract
            text = textract.process(file_path).decode('utf-8')
            file_type = "DOC"
        else:
            raise ValueError(f"Format de fichier non pris en charge: {extension}")
        
        logger.info(f"✅ {file_type} chargé: {file_path}")
        return [Document(page_content=text, metadata={"source": file_path})]

#CSV
class CSVLoader(DocumentLoader):
    def load(self, file_path: str) -> List[Document]:
        if not os.path.exists(file_path):
            raise FileNotFoundError(f" Le fichier {file_path} n'existe pas.")

        df = pd.read_csv(file_path)
        rows = df.astype(str).values.tolist()
        documents = []
        
        for row in rows:
            text = " | ".join(row)
            documents.append(Document(page_content=text, metadata={"source": file_path}))
        #row_texts = [" | ".join(row) for row in rows]

        logger.info(f"✅ CSV chargé: {file_path}")
        return documents

#TXT
class TXTLoader(DocumentLoader):
    def load(self, file_path: str) -> List[Document]:
        if not os.path.exists(file_path):
            raise FileNotFoundError(f" Le fichier {file_path} n'existe pas.")

        with open(file_path, 'r', encoding='utf-8') as file:
            text = file.read()

        logger.info(f"✅ TXT chargé: {file_path}")
        return [Document(page_content=text, metadata={"source": file_path})]
    
# RTFloader classes
class RTFLoader(DocumentLoader):
    def load(self, file_path: str) -> List[Document]:
        if not os.path.exists(file_path):
            raise FileNotFoundError(f"Le fichier {file_path} n'existe pas.")
            
        try:
            with open(file_path, 'r', encoding='utf-8') as file:
                rtf_content = file.read()
                text = rtf_to_text(rtf_content)
                
            logger.info(f"✅ RTF chargé: {file_path}")
            return [Document(
                page_content=text,
                metadata={
                    "source": file_path,
                    "format": "RTF"
                }
            )]
        except Exception as e:
            logger.error(f"❌ Erreur lors du chargement du RTF {file_path}: {str(e)}")
            # Try fallback method with textract if striprtf fails
            try:
                text = textract.process(file_path).decode('utf-8')
                logger.info(f"✅ RTF chargé (via textract): {file_path}")
                return [Document(
                    page_content=text,
                    metadata={
                        "source": file_path,
                        "format": "RTF",
                        "fallback": "textract"
                    }
                )]
            except Exception as e2:
                raise Exception(f"Failed to load RTF file using both methods: {str(e)} | {str(e2)}")

class MarkdownLoader(DocumentLoader):
    def load(self, file_path: str) -> List[Document]:
        with open(file_path, 'r', encoding='utf-8') as f:
            text = f.read()
        return [Document(page_content=text, metadata={"source": file_path})]

class PowerPointLoader(DocumentLoader):
    def load(self, file_path: str) -> List[Document]:
        text = textract.process(file_path).decode('utf-8')
        return [Document(page_content=text, metadata={"source": file_path})]

class XMLLoader(DocumentLoader):
    def load(self, file_path: str) -> List[Document]:
        tree = ET.parse(file_path)
        root = tree.getroot()
        text = ET.tostring(root, encoding='unicode', method='xml')
        return [Document(page_content=text, metadata={"source": file_path})]

class YAMLLoader(DocumentLoader):
    def load(self, file_path: str) -> List[Document]:
        with open(file_path, 'r') as f:
            data = yaml.safe_load(f)
        text = yaml.dump(data, allow_unicode=True)
        return [Document(page_content=text, metadata={"source": file_path})]

class CADLoader(DocumentLoader):
    def load(self, file_path: str) -> List[Document]:
        doc = ezdxf.readfile(file_path)
        text = str(doc.header) + "\n"
        for entity in doc.modelspace():
            text += str(entity) + "\n"
        return [Document(page_content=text, metadata={"source": file_path})]

class STLLoader(DocumentLoader):
    def load(self, file_path: str) -> List[Document]:
        try:
            # Load the STL file
            stl_mesh = numpy_stl.Mesh.from_file(file_path)
            
            # Extract basic geometry information
            vertices = len(stl_mesh.vectors)
            volume, cog, inertia = stl_mesh.get_mass_properties()
            
            # Create a structured text representation
            text = f"""STL File Analysis:
Number of vertices: {vertices}
Volume: {volume:.2f} cubic units
Center of gravity: [{', '.join(f'{x:.2f}' for x in cog)}]
"""
            
            logger.info(f"✅ STL chargé: {file_path}")
            return [Document(
                page_content=text,
                metadata={
                    "source": file_path,
                    "format": "STL",
                    "vertices": vertices,
                    "volume": float(volume)
                }
            )]
            
        except Exception as e:
            logger.error(f"❌ Erreur lors du chargement du STL {file_path}: {str(e)}")
            raise

# TODO: Implement these loaders
class DatabaseLoader(DocumentLoader):
    """For handling .db files using sqlite3"""
    pass

class ConfigLoader(DocumentLoader):
    """For handling .ini and .conf files using configparser"""
    pass

class VectorImageLoader(DocumentLoader):
    """For handling .svg and .eps files using svglib and renderPM"""
    pass

# --------------------------------------------------------------------
# STEP 3: Factory/File Extension
# --------------------------------------------------------------------

class DocumentLoaderFactory:

    @staticmethod
    def get_loader(file_path: str) -> DocumentLoader:
        if not os.path.exists(file_path):
            raise FileNotFoundError(f"Le fichier {file_path} n'existe pas.")

        extension = Path(file_path).suffix.lower()

        # Documents
        if extension in SUPPORTED_FILE_TYPES['documents']:
            if extension == '.pdf':
                return PDFLoader()
            elif extension in ['.doc', '.docx', '.odt']:
                return WordDocumentLoader()
            elif extension in ['.ppt', '.pptx', '.odp']:
                return PowerPointLoader()
            elif extension == '.rtf':
                return RTFLoader()
            elif extension in ['.md', '.log', '.txt']:
                return TXTLoader()
            elif extension in ['.xlsx', '.xls', '.ods']:
                return ExcelLoader()

        # Data files
        elif extension in SUPPORTED_FILE_TYPES['data_files']:
            if extension == '.csv':
                return CSVLoader()
            elif extension == '.json':
                return JSONLoader()
            elif extension == '.xml':
                return XMLLoader()
            elif extension in ['.yaml', '.yml']:
                return YAMLLoader()
            elif extension in ['.ini', '.conf']:
                return ConfigLoader()
            elif extension == '.db':
                return DatabaseLoader()

        # Technical files
        elif extension in SUPPORTED_FILE_TYPES['technical']:
            if extension in ['.dxf', '.dwg']:
                return CADLoader()
            elif extension in ['.stl', '.obj']:
                return STLLoader()

        # Images
        elif extension in SUPPORTED_FILE_TYPES['images']:
            if extension in ['.png', '.jpg', '.jpeg', '.tiff', '.bmp', '.tif', '.webp']:
                return ImageLoader()
            elif extension in ['.svg', '.eps']:
                return VectorImageLoader()

        raise ValueError(
            f"❌ Le format {extension} n'est pas supporté. "
            f"Formats supportés: {', '.join(sorted(ALL_SUPPORTED_EXTENSIONS))}"
        )
# --------------------------------------------------------------------
# STEP 4: Helper function
# --------------------------------------------------------------------

def load_document(file_path: str) -> List[Document]:
    
    """
    Loads a document from a given file path using the appropriate loader.

    Args:
        file_path (str): The path to the file to be loaded.

    Returns:
        Union[str, List[str]]: The extracted content from the file.
    """
    
    loader = DocumentLoaderFactory.get_loader(file_path)
    return loader.load(file_path)

# --------------------------------------------------------------------
# STEP 5: Load All Files from a Directory
# --------------------------------------------------------------------
def load_all_documents_from_directory(directory: str) -> List[Document]:
    if not os.path.exists(directory):
        raise FileNotFoundError(f"Le dossier {directory} n'existe pas.")

    all_documents = []
    for filename in os.listdir(directory):
        file_path = os.path.join(directory, filename)

        if os.path.isfile(file_path):
            try:
                documents = load_document(file_path)
                all_documents.extend(documents)
            except Exception as e:
                logger.warning(f"⚠️ Erreur lors du chargement de {file_path}: {e}")

    logger.info(f"Total documents chargés: {len(all_documents)}")
    return all_documents

# --------------------------------------------------------------------
# STEP 6: Exemple d’utilisation directe (facultatif)
# --------------------------------------------------------------------

if __name__ == "__main__":
    
    try: 
        directory_path = os.path.join(DATA_DIR, "raw")
        all_docs = load_all_documents_from_directory(directory_path)
        
        for idx, doc in enumerate(all_docs):
            print(f"📄 Document {idx + 1}")
            print(doc.page_content[:300])  # Just previewing the first 300 characters
            print("-" * 80)
            
    except Exception as e:
        logger.error(f"⚠️ Erreur lors du chargement des documents : {e}")


# if __name__ == "__main__":
#     try:
#         content = load_document(os.path.join(DATA_DIR, "TestFile.txt"))
#         for doc in content:
#             print(doc.page_content[:300])
#             print("-" * 50)
#     except Exception as e:
#         logger.error(f"⚠️ Erreur: {e}")